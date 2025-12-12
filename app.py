from flask import Flask, render_template, request, make_response
try:
    from weasyprint import HTML
    WEASYPRINT_OK = True
    WEASYPRINT_ERROR = None
except Exception as exc:  # ambiente Windows pode faltar libs do sistema
    HTML = None
    WEASYPRINT_OK = False
    WEASYPRINT_ERROR = exc
import io
import re
import json
from typing import Dict, Any

import requests
from bs4 import BeautifulSoup
from docx import Document

app = Flask(__name__)


def sanitize_filename(name: str) -> str:
    name = name.strip().lower()
    name = re.sub(r"[^a-z0-9]+", "_", name)
    name = re.sub(r"_+", "_", name)
    return name.strip("_") or "curriculo"


def normalize_text(text: str) -> str:
    """Limpeza simples: tira espaços extras e garante início com maiúscula."""
    text = (text or "").strip()
    if not text:
        return ""
    # Espaços internos
    text = re.sub(r"\s+", " ", text)
    # Primeira letra maiúscula
    return text[0].upper() + text[1:]


def normalize_period(periodo: str) -> str:
    """Padroniza o separador do período em ' — ' sem tentar parser complexo."""
    if not periodo:
        return ""
    periodo = re.sub(r"\s*[-–—]+\s*", " — ", periodo.strip())
    return periodo


def fetch_job_keywords(url: str, max_keywords: int = 30) -> set:
    """Busca vaga por URL e extrai palavras-chave simples."""
    if not url:
        return set()
    try:
        resp = requests.get(url, timeout=5)
        resp.raise_for_status()
    except Exception:
        return set()

    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text(" ")
    text = re.sub(r"\s+", " ", text)
    words = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]{4,}", text.lower())
    freq: Dict[str, int] = {}
    for w in words:
        freq[w] = freq.get(w, 0) + 1
    sorted_words = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return {w for w, _ in sorted_words[:max_keywords]}


def score_text(text: str, keywords: set) -> int:
    if not text or not keywords:
        return 0
    words = set(re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]{4,}", text.lower()))
    return len(words & keywords)


def generate_word(cv_data: Dict[str, Any]) -> io.BytesIO:
    """Gera um DOCX simples a partir dos dados do currículo."""
    doc = Document()

    doc.add_heading(cv_data.get("nome") or "Currículo", level=0)
    if cv_data.get("titulo"):
        doc.add_paragraph(cv_data["titulo"])

    contato_parts = []
    for k in ("email", "telefone", "endereco", "portfolio"):
        v = cv_data.get(k)
        if v:
            contato_parts.append(str(v))
    if contato_parts:
        doc.add_paragraph(" | ".join(contato_parts))

    if cv_data.get("resumo"):
        doc.add_heading("Resumo profissional", level=1)
        doc.add_paragraph(cv_data["resumo"])

    if cv_data.get("experiencias"):
        doc.add_heading("Experiência profissional", level=1)
        for exp in cv_data["experiencias"]:
            titulo = " | ".join(
                [
                    part
                    for part in [exp.get("cargo"), exp.get("empresa"), exp.get("periodo")]
                    if part
                ]
            )
            if titulo:
                doc.add_paragraph(titulo, style="List Bullet")
            if exp.get("descricao"):
                doc.add_paragraph(exp["descricao"])

    if cv_data.get("formacoes"):
        doc.add_heading("Formação", level=1)
        for edu in cv_data["formacoes"]:
            linha = " | ".join(
                [
                    part
                    for part in [
                        edu.get("curso"),
                        edu.get("instituicao"),
                        edu.get("cidade"),
                        edu.get("ano"),
                        edu.get("status"),
                    ]
                    if part
                ]
            )
            if linha:
                doc.add_paragraph(linha, style="List Bullet")

    if any(
        cv_data.get(k)
        for k in ("skills_tecnicas", "skills_comportamentais", "skills_outras")
    ):
        doc.add_heading("Habilidades", level=1)
        for label, key in (
            ("Técnicas", "skills_tecnicas"),
            ("Comportamentais", "skills_comportamentais"),
            ("Outras", "skills_outras"),
        ):
            skills = cv_data.get(key) or []
            if skills:
                doc.add_paragraph(f"{label}: {', '.join(skills)}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")


@app.route("/gerar", methods=["POST"])
def gerar():
    if not WEASYPRINT_OK:
        # Falha amigável se as bibliotecas de sistema do WeasyPrint não estiverem instaladas
        msg = (
            "Não foi possível gerar o PDF neste ambiente porque o WeasyPrint "
            "não conseguiu carregar as bibliotecas de sistema necessárias. "
            "Em Windows, instale o pacote oficial conforme a documentação: "
            "https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation.\n\n"
        )
        msg += f"Detalhes técnicos: {WEASYPRINT_ERROR}"
        return msg, 500

    # Limites simples de segurança (poderia ser ajustado)
    max_len = 4000

    def limited(field):
        value = request.form.get(field, "").strip()
        return value[:max_len]

    def limited_list(field):
        values = request.form.getlist(field)
        result = []
        for v in values:
            v = (v or "").strip()
            result.append(v[:max_len])
        return result

    nome = normalize_text(limited("nome"))
    titulo = normalize_text(limited("titulo"))
    email = limited("email")
    telefone = limited("telefone")
    endereco = normalize_text(limited("endereco"))
    portfolio = limited("link_portfolio")
    # Foto: prioriza upload, cai para URL se não houver arquivo
    foto_url = limited("foto_url")
    foto_arquivo = request.files.get("foto_arquivo")
    if foto_arquivo and foto_arquivo.filename:
        # WeasyPrint aceita data URL; mantemos tudo em memória, sem salvar em disco
        foto_bytes = foto_arquivo.read()
        import base64

        mime = foto_arquivo.mimetype or "image/jpeg"
        b64 = base64.b64encode(foto_bytes).decode("ascii")
        foto_url = f"data:{mime};base64,{b64}"
    resumo = normalize_text(limited("resumo"))[:1000]

    job_url = request.form.get("job_url", "").strip()
    output_format = request.form.get("output_format", "pdf").lower()
    template_style = request.form.get("template_style", "corporativo").lower()

    # Experiências profissionais (múltiplas)
    exp_empresas = limited_list("exp_empresa")
    exp_cargos = limited_list("exp_cargo")
    exp_periodos = [normalize_period(p) for p in limited_list("exp_periodo")]
    exp_descricoes = limited_list("exp_descricao")
    exp_locais = limited_list("exp_local")
    exp_conquistas = limited_list("exp_conquistas")
    exp_tecnologias = limited_list("exp_tech")

    experiencias = []
    for empresa, cargo, periodo, descricao, local, conquistas, tecnologias in zip(
        exp_empresas,
        exp_cargos,
        exp_periodos,
        exp_descricoes,
        exp_locais,
        exp_conquistas,
        exp_tecnologias,
    ):
        if not (empresa or cargo or periodo or descricao or local or conquistas or tecnologias):
            continue
        experiencias.append(
            {
                "empresa": empresa,
                "cargo": cargo,
                "periodo": periodo,
                "descricao": descricao,
                "local": local,
                "conquistas": conquistas,
                "tecnologias": tecnologias,
            }
        )

    # Formação / educação (múltiplas)
    edu_cursos = limited_list("edu_curso")
    edu_insts = limited_list("edu_inst")
    edu_cidades = limited_list("edu_cidade")
    edu_anos = limited_list("edu_ano")
    edu_status_list = limited_list("edu_status")

    formacoes = []
    for curso, inst, cidade, ano, status in zip(
        edu_cursos, edu_insts, edu_cidades, edu_anos, edu_status_list
    ):
        if not (curso or inst or cidade or ano or status):
            continue
        formacoes.append(
            {
                "curso": curso,
                "instituicao": inst,
                "cidade": cidade,
                "ano": ano,
                "status": status,
            }
        )

    # Habilidades em categorias
    skills_tecnicas_raw = limited("skills_tecnicas")
    skills_comport_raw = limited("skills_comportamentais")
    skills_outras_raw = limited("skills_outras")

    def parse_skills(raw: str):
        return [s.strip() for s in raw.split(",") if s.strip()]

    skills_tecnicas = parse_skills(skills_tecnicas_raw)
    skills_comportamentais = parse_skills(skills_comport_raw)
    skills_outras = parse_skills(skills_outras_raw)

    # Certificações
    cert_nomes = limited_list("cert_nome")
    cert_insts = limited_list("cert_inst")
    cert_anos = limited_list("cert_ano")
    cert_codigos = limited_list("cert_codigo")

    certificacoes = []
    for nome_c, inst_c, ano_c, cod in zip(
        cert_nomes, cert_insts, cert_anos, cert_codigos
    ):
        if not (nome_c or inst_c or ano_c or cod):
            continue
        certificacoes.append(
            {
                "nome": nome_c,
                "instituicao": inst_c,
                "ano": ano_c,
                "codigo": cod,
            }
        )

    # Projetos
    proj_nomes = limited_list("proj_nome")
    proj_tecs = limited_list("proj_tec")
    proj_descs = limited_list("proj_desc")
    proj_links = limited_list("proj_link")

    projetos = []
    for nome_p, tec_p, desc_p, link_p in zip(
        proj_nomes, proj_tecs, proj_descs, proj_links
    ):
        if not (nome_p or tec_p or desc_p or link_p):
            continue
        projetos.append(
            {
                "nome": nome_p,
                "tecnologias": tec_p,
                "descricao": desc_p,
                "link": link_p,
            }
        )

    # Idiomas
    idioma_nomes = limited_list("idioma_nome")
    idioma_niveis = limited_list("idioma_nivel")

    idiomas = []
    for nome_i, nivel_i in zip(idioma_nomes, idioma_niveis):
        if not (nome_i or nivel_i):
            continue
        idiomas.append({"nome": nome_i, "nivel": nivel_i})

    # Cursos / workshops
    curso_extra_nomes = limited_list("curso_extra_nome")
    curso_extra_cargas = limited_list("curso_extra_carga")
    curso_extra_insts = limited_list("curso_extra_inst")
    curso_extra_anos = limited_list("curso_extra_ano")

    cursos_extra = []
    for nome_e, carga_e, inst_e, ano_e in zip(
        curso_extra_nomes, curso_extra_cargas, curso_extra_insts, curso_extra_anos
    ):
        if not (nome_e or carga_e or inst_e or ano_e):
            continue
        cursos_extra.append(
            {
                "nome": nome_e,
                "carga": carga_e,
                "instituicao": inst_e,
                "ano": ano_e,
            }
        )

    # Prêmios e reconhecimentos
    premio_titulos = limited_list("premio_titulo")
    premio_insts = limited_list("premio_inst")
    premio_anos = limited_list("premio_ano")
    premio_descs = limited_list("premio_desc")

    premios = []
    for titulo_p, inst_p, ano_p, desc_pr in zip(
        premio_titulos, premio_insts, premio_anos, premio_descs
    ):
        if not (titulo_p or inst_p or ano_p or desc_pr):
            continue
        premios.append(
            {
                "titulo": titulo_p,
                "instituicao": inst_p,
                "ano": ano_p,
                "descricao": desc_pr,
            }
        )

    # Voluntariado
    vol_orgs = limited_list("vol_org")
    vol_funcoes = limited_list("vol_funcao")
    vol_periodos = limited_list("vol_periodo")
    vol_descs = limited_list("vol_desc")

    voluntariados = []
    for org_v, func_v, per_v, desc_v in zip(
        vol_orgs, vol_funcoes, vol_periodos, vol_descs
    ):
        if not (org_v or func_v or per_v or desc_v):
            continue
        voluntariados.append(
            {
                "organizacao": org_v,
                "funcao": func_v,
                "periodo": per_v,
                "descricao": desc_v,
            }
        )

    # Publicações (texto livre, opcional)
    publicacoes_texto = limited("publicacoes")

    # Inteligência básica: usar vaga (se URL fornecida) para priorizar palavras-chave
    keywords = fetch_job_keywords(job_url)
    if keywords:
        for lista_nome in ("skills_tecnicas", "skills_comportamentais", "skills_outras"):
            lst = locals().get(lista_nome) or []
            lst.sort(key=lambda s: score_text(s, keywords), reverse=True)
            locals()[lista_nome] = lst

        experiencias.sort(
            key=lambda e: score_text(
                " ".join(
                    [
                        e.get("cargo", ""),
                        e.get("empresa", ""),
                        e.get("descricao", ""),
                        e.get("conquistas", ""),
                        e.get("tecnologias", ""),
                    ]
                ),
                keywords,
            ),
            reverse=True,
        )

    # Pacote único de dados do currículo para PDF/Word/JSON
    cv_data: Dict[str, Any] = {
        "nome": nome,
        "titulo": titulo,
        "email": email,
        "telefone": telefone,
        "endereco": endereco,
        "portfolio": portfolio,
        "foto_url": foto_url,
        "resumo": resumo,
        "experiencias": experiencias,
        "formacoes": formacoes,
        "skills_tecnicas": skills_tecnicas,
        "skills_comportamentais": skills_comportamentais,
        "skills_outras": skills_outras,
        "certificacoes": certificacoes,
        "projetos": projetos,
        "idiomas": idiomas,
        "cursos_extra": cursos_extra,
        "premios": premios,
        "voluntariados": voluntariados,
        "publicacoes_texto": publicacoes_texto,
    }

    # Exportar em JSON
    if output_format == "json":
        payload = json.dumps(cv_data, ensure_ascii=False, indent=2)
        safe_name = sanitize_filename(nome or "curriculo")
        filename = f"curriculo_{safe_name}.json"
        resp = make_response(payload)
        resp.headers["Content-Type"] = "application/json; charset=utf-8"
        resp.headers["Content-Disposition"] = f"attachment; filename={filename}"
        return resp

    # Exportar em Word
    if output_format == "word":
        doc_io = generate_word(cv_data)
        safe_name = sanitize_filename(nome or "curriculo")
        filename = f"curriculo_{safe_name}.docx"
        response = make_response(doc_io.read())
        response.headers[
            "Content-Type"
        ] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        response.headers["Content-Disposition"] = f"attachment; filename={filename}"
        return response

    # Escolha de template PDF
    if template_style == "minimalista":
        template_pdf = "resume_template_minimal.html"
    elif template_style == "ats":
        template_pdf = "resume_template_ats.html"
    else:
        template_pdf = "resume_template.html"

    html = render_template(template_pdf, **cv_data)

    pdf_io = io.BytesIO()
    try:
        HTML(string=html, base_url=request.base_url).write_pdf(pdf_io)
    except Exception as exc:  # fallback amigável caso o WeasyPrint falhe no ambiente
        message = (
            "Erro ao gerar PDF. Parece que o WeasyPrint não está totalmente "
            "configurado neste sistema (bibliotecas gráficas ausentes). "
            "Siga as instruções de instalação em https://weasyprint.org/ "
            "ou use a imagem Docker incluída no projeto.\n\n"
            f"Detalhes técnicos: {type(exc).__name__}: {exc}"
        )
        response = make_response(message, 500)
        response.headers["Content-Type"] = "text/plain; charset=utf-8"
        return response

    pdf_io.seek(0)

    safe_name = sanitize_filename(nome or "curriculo")
    filename = f"curriculo_{safe_name}.pdf"

    response = make_response(pdf_io.read())
    response.headers["Content-Type"] = "application/pdf"
    response.headers["Content-Disposition"] = f"attachment; filename={filename}"

    return response


if __name__ == "__main__":
    app.run(debug=True)
